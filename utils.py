from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def save_to_msword(text):
    doc = Document()

    # Markdown content returned from the LLM
    report_text = text

    # Split lines in markdown
    lines = report_text.splitlines()

    # Detect hyperlink format: [text](url)
    hyperlink_pattern = re.compile(r'\[(.*?)\]\((https?://\S+)\)')

    for line in lines:
        line = line.strip() 

        # Check title level based on nums of #
        if line.startswith("##### "):  # level 5
            doc.add_heading(line[6:], level=4)
        elif line.startswith("#### "):  # level 4
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):  # level 3
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):  # level 2
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):  # level 1
            doc.add_heading(line[2:], level=1)
        elif line:  # text or hyperlinks
            paragraph = doc.add_paragraph()
            
            # Find all hyperlinks
            pos = 0
            for match in hyperlink_pattern.finditer(line):
                # Add text before hyperlink
                paragraph.add_run(line[pos:match.start()])
                
                # Add hyperlink
                text, url = match.groups()
                add_hyperlink(paragraph, url, '['+text+']')
                
                # Move to next position
                pos = match.end()
            
            # Add the rest of content after hyperlink
            paragraph.add_run(line[pos:])

    # Save
    doc.save("article.docx")

def add_hyperlink(paragraph, url, text):
    """
    Save hyperlink format into MS Word with blue color.
    """
    # Create an object for hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Setting color and underline for hyperlink
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    
    # Blue for hyperlink
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Màu xanh lam
    rPr.append(color)
    
    # # Underline hyperlink
    # underline = OxmlElement('w:u')
    # underline.set(qn('w:val'), 'single')
    # rPr.append(underline)

    new_run.append(rPr)
    text_run = OxmlElement('w:t')
    text_run.text = text
    new_run.append(text_run)
    hyperlink.append(new_run)

    # Add hyperlink
    paragraph._p.append(hyperlink)